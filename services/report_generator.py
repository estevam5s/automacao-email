import json
import csv
import io
from typing import List
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from xml.etree.ElementTree import Element, SubElement, tostring
from xml.dom import minidom
from data.models.funcionario import Funcionario
from config.settings import settings

class ReportGenerator:
    def __init__(self, funcionarios: List[Funcionario], dia_trabalho: date):
        self.funcionarios = funcionarios
        self.dia_trabalho = dia_trabalho
        self.dia_semana = settings.DIAS_SEMANA.get(dia_trabalho.weekday(), "")
        self.total = sum(f.valor_10_percent for f in funcionarios)

    def _get_table_html(self) -> str:
        rows = ""
        for f in self.funcionarios:
            rows += f"""
            <tr>
                <td>{f.nome}</td>
                <td>R$ {f.valor_10_percent:.2f}</td>
                <td>{f.hora_entrada}</td>
                <td>{f.hora_saida}</td>
                <td>{f.observacao or '-'}</td>
            </tr>"""
        return rows

    def _format_xml(self, xml_str: str) -> str:
        dom = minidom.parseString(xml_str)
        return dom.toprettyxml(indent="  ")

    def generate_docx(self) -> bytes:
        doc = Document()
        
        title = doc.add_heading('RELATÓRIO DE SALÁRIOS DOS GARÇONS', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = doc.add_paragraph()
        info.add_run(f"Data: {self.dia_trabalho.strftime('%d/%m/%Y')}\n").bold = True
        info.add_run(f"Dia da Semana: {self.dia_semana}\n").bold = True
        info.add_run(f"Total de Funcionários: {len(self.funcionarios)}\n").bold = True
        info.add_run(f"Total a Pagar: R$ {self.total:.2f}").bold = True
        
        table = doc.add_table(rows=1, cols=5, style='Table Grid')
        headers = ['Nome', '10% (R$)', 'Entrada', 'Saída', 'Observação']
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for f in self.funcionarios:
            row_cells = table.add_row().cells
            row_cells[0].text = f.nome
            row_cells[1].text = f"{f.valor_10_percent:.2f}"
            row_cells[2].text = f.hora_entrada
            row_cells[3].text = f.hora_saida
            row_cells[4].text = f.observacao or "-"
        
        doc.add_paragraph(f"\nGerado em: {date.today().strftime('%d/%m/%Y às %H:%M')}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_excel(self) -> bytes:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Relatório"
        
        ws.merge_cells('A1:E1')
        ws['A1'] = 'RELATÓRIO DE SALÁRIOS DOS GARÇONS'
        ws['A1'].font = Font(size=14, bold=True, color="FFFFFF")
        ws['A1'].fill = PatternFill("solid", fgColor="2E7D32")
        ws['A1'].alignment = Alignment(horizontal="center")
        
        ws['A2'] = f"Data: {self.dia_trabalho.strftime('%d/%m/%Y')}"
        ws['B2'] = f"Dia: {self.dia_semana}"
        ws['A3'] = f"Total Funcionários: {len(self.funcionarios)}"
        ws['B3'] = f"Total a Pagar: R$ {self.total:.2f}"
        
        headers = ['Nome', '10% (R$)', 'Entrada', 'Saída', 'Observação']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=5, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1565C0")
            cell.alignment = Alignment(horizontal="center")
        
        for row, f in enumerate(self.funcionarios, 6):
            ws.cell(row=row, column=1, value=f.nome)
            ws.cell(row=row, column=2, value=f.valor_10_percent)
            ws.cell(row=row, column=3, value=f.hora_entrada)
            ws.cell(row=row, column=4, value=f.hora_saida)
            ws.cell(row=row, column=5, value=f.observacao or "-")
        
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 12
        ws.column_dimensions['E'].width = 30
        
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def generate_csv(self) -> str:
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(['Nome', '10% (R$)', 'Entrada', 'Saída', 'Observação'])
        
        for f in self.funcionarios:
            writer.writerow([
                f.nome,
                f"{f.valor_10_percent:.2f}",
                f.hora_entrada,
                f.hora_saida,
                f.observacao or ""
            ])
        
        return output.getvalue()

    def generate_json(self) -> str:
        data = {
            "relatorio": {
                "data": self.dia_trabalho.isoformat(),
                "dia_semana": self.dia_semana,
                "total_funcionarios": len(self.funcionarios),
                "total_valores": self.total
            },
            "funcionarios": [
                {
                    "nome": f.nome,
                    "valor_10_percent": f.valor_10_percent,
                    "hora_entrada": f.hora_entrada,
                    "hora_saida": f.hora_saida,
                    "observacao": f.observacao
                }
                for f in self.funcionarios
            ]
        }
        return json.dumps(data, ensure_ascii=False, indent=2)

    def generate_xml(self) -> str:
        root = Element('relatorio')
        root.set('data', self.dia_trabalho.isoformat())
        root.set('dia_semana', self.dia_semana)
        
        meta = SubElement(root, 'resumo')
        SubElement(meta, 'total_funcionarios').text = str(len(self.funcionarios))
        SubElement(meta, 'total_valores').text = f"{self.total:.2f}"
        
        funcs = SubElement(root, 'funcionarios')
        for f in self.funcionarios:
            func = SubElement(funcs, 'funcionario')
            SubElement(func, 'nome').text = f.nome
            SubElement(func, 'valor_10_percent').text = f"{f.valor_10_percent:.2f}"
            SubElement(func, 'hora_entrada').text = f.hora_entrada
            SubElement(func, 'hora_saida').text = f.hora_saida
            SubElement(func, 'observacao').text = f.observacao or ""
        
        return self._format_xml(tostring(root, encoding='unicode'))

    def generate_html(self) -> str:
        return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
    <meta charset="UTF-8">
    <title>Relatório de Salários - {self.dia_trabalho.strftime('%d/%m/%Y')}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        h1 {{ color: #2E7D32; text-align: center; }}
        .info {{ background: #E8F5E9; padding: 15px; border-radius: 8px; margin: 20px 0; }}
        .info p {{ margin: 5px 0; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
        th {{ background: #1565C0; color: white; padding: 12px; text-align: left; }}
        td {{ padding: 10px; border-bottom: 1px solid #ddd; }}
        tr:nth-child(even) {{ background: #f9f9f9; }}
        .total {{ font-weight: bold; font-size: 18px; color: #2E7D32; }}
        .footer {{ margin-top: 30px; text-align: center; color: #666; font-size: 12px; }}
    </style>
</head>
<body>
    <h1>RELATÓRIO DE SALÁRIOS DOS GARÇONS</h1>
    <div class="info">
        <p><strong>Data:</strong> {self.dia_trabalho.strftime('%d/%m/%Y')}</p>
        <p><strong>Dia da Semana:</strong> {self.dia_semana}</p>
        <p><strong>Total de Funcionários:</strong> {len(self.funcionarios)}</p>
        <p class="total">Total a Pagar: R$ {self.total:.2f}</p>
    </div>
    <table>
        <thead>
            <tr>
                <th>Nome</th>
                <th>10% (R$)</th>
                <th>Entrada</th>
                <th>Saída</th>
                <th>Observação</th>
            </tr>
        </thead>
        <tbody>
            {self._get_table_html()}
        </tbody>
    </table>
    <div class="footer">
        Gerado em {date.today().strftime('%d/%m/%Y às %H:%M')}
    </div>
</body>
</html>"""

    def generate_all(self) -> dict:
        return {
            "docx": self.generate_docx(),
            "excel": self.generate_excel(),
            "csv": self.generate_csv().encode('utf-8'),
            "json": self.generate_json().encode('utf-8'),
            "xml": self.generate_xml().encode('utf-8'),
            "html": self.generate_html().encode('utf-8')
        }
